"""Explainable, privacy-friendly CV analysis.

The checker intentionally uses transparent heuristics instead of pretending that
an opaque score can predict a hiring decision.  It never sends CV content to an
external service.
"""

from __future__ import annotations

from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
import re
from typing import BinaryIO, Iterable
from zipfile import BadZipFile, ZipFile


ANALYZER_VERSION = "2.0"
MAX_TEXT_CHARACTERS = 150_000


class CVExtractionError(ValueError):
    """Raised when a CV cannot be read safely enough to analyse."""


class UnsupportedCVFormat(CVExtractionError):
    """Raised when the uploaded format has no reliable parser."""


@dataclass
class ExtractionResult:
    text: str
    file_type: str
    page_count: int | None = None
    table_count: int | None = None
    warnings: list[str] = field(default_factory=list)


SECTION_ALIASES = {
    "summary": (
        "summary",
        "profile",
        "professional profile",
        "professional summary",
        "career objective",
        "objective",
        "about me",
    ),
    "experience": (
        "experience",
        "work experience",
        "professional experience",
        "employment history",
        "career history",
    ),
    "education": (
        "education",
        "academic background",
        "academic qualifications",
        "qualifications",
    ),
    "skills": (
        "skills",
        "technical skills",
        "core competencies",
        "competencies",
        "expertise",
        "tools",
    ),
    "projects": ("projects", "selected projects", "academic projects", "portfolio"),
    "certifications": (
        "certifications",
        "certificates",
        "professional development",
        "training",
    ),
}

ACTION_VERBS = {
    "achieved",
    "analysed",
    "analyzed",
    "automated",
    "built",
    "collaborated",
    "created",
    "delivered",
    "designed",
    "developed",
    "drove",
    "established",
    "executed",
    "generated",
    "grew",
    "implemented",
    "improved",
    "increased",
    "launched",
    "led",
    "managed",
    "mentored",
    "negotiated",
    "optimized",
    "organised",
    "organized",
    "planned",
    "presented",
    "produced",
    "reduced",
    "researched",
    "resolved",
    "saved",
    "secured",
    "streamlined",
    "supported",
    "tested",
    "trained",
    "transformed",
    "won",
}

GENERIC_PHRASES = {
    "hard worker",
    "hardworking",
    "results-driven",
    "results driven",
    "team player",
    "detail-oriented",
    "detail oriented",
    "go-getter",
    "self-motivated",
    "self motivated",
    "excellent communication skills",
    "works well under pressure",
}

SKILL_TERMS = {
    "adobe creative suite",
    "agile",
    "amazon web services",
    "angular",
    "autocad",
    "aws",
    "azure",
    "business analysis",
    "c#",
    "c++",
    "canva",
    "cloud computing",
    "communication",
    "content marketing",
    "copywriting",
    "css",
    "customer service",
    "cybersecurity",
    "data analysis",
    "data visualization",
    "django",
    "docker",
    "excel",
    "figma",
    "financial analysis",
    "flask",
    "git",
    "google analytics",
    "graphic design",
    "html",
    "java",
    "javascript",
    "jira",
    "kubernetes",
    "leadership",
    "machine learning",
    "marketing",
    "microsoft office",
    "mongodb",
    "mysql",
    "node.js",
    "notion",
    "numpy",
    "pandas",
    "photoshop",
    "php",
    "postgresql",
    "power bi",
    "presentation",
    "project management",
    "public speaking",
    "python",
    "react",
    "research",
    "sales",
    "seo",
    "social media",
    "sql",
    "tableau",
    "tensorflow",
    "typescript",
    "ui design",
    "user research",
    "ux design",
    "wordpress",
}

STOP_WORDS = {
    "about",
    "after",
    "also",
    "and",
    "are",
    "because",
    "been",
    "being",
    "but",
    "can",
    "candidate",
    "company",
    "for",
    "from",
    "have",
    "into",
    "job",
    "must",
    "our",
    "role",
    "should",
    "that",
    "the",
    "their",
    "them",
    "they",
    "this",
    "through",
    "using",
    "will",
    "with",
    "work",
    "working",
    "years",
    "you",
    "your",
}


def extract_cv(file_obj: BinaryIO, filename: str) -> ExtractionResult:
    """Extract text and limited structural metadata from a PDF or DOCX file."""

    suffix = Path(filename).suffix.lower()
    try:
        original_position = file_obj.tell()
        file_obj.seek(0, 2)
        file_size = file_obj.tell()
        file_obj.seek(original_position)
    except (AttributeError, OSError):
        file_size = getattr(file_obj, "size", None)
    if file_size is not None and file_size > 10 * 1024 * 1024:
        raise CVExtractionError(
            "This file is too large to analyse safely. Upload a CV under 10 MB."
        )
    try:
        file_obj.seek(0)
    except (AttributeError, OSError):
        pass

    if suffix == ".pdf":
        return _extract_pdf(file_obj)
    if suffix == ".docx":
        return _extract_docx(file_obj)
    if suffix == ".doc":
        raise UnsupportedCVFormat(
            "Legacy .doc files cannot be analysed reliably. Save the CV as PDF "
            "or DOCX and upload the converted copy."
        )
    raise UnsupportedCVFormat("Only PDF and DOCX files can be analysed.")


def _extract_pdf(file_obj: BinaryIO) -> ExtractionResult:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - depends on deployment
        raise CVExtractionError(
            "PDF analysis is temporarily unavailable because the PDF parser is "
            "not installed."
        ) from exc

    try:
        reader = PdfReader(file_obj)
        if reader.is_encrypted:
            try:
                unlocked = reader.decrypt("")
            except Exception as exc:
                raise CVExtractionError(
                    "This PDF is password-protected. Upload an unlocked copy."
                ) from exc
            if not unlocked:
                raise CVExtractionError(
                    "This PDF is password-protected. Upload an unlocked copy."
                )
        if len(reader.pages) > 50:
            raise CVExtractionError(
                "This PDF has too many pages for a CV review. Upload a focused CV "
                "of 50 pages or fewer."
            )

        pages = []
        warnings = []
        for page_number, page in enumerate(reader.pages, start=1):
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                pages.append("")
                warnings.append(f"Text could not be read from page {page_number}.")
    except CVExtractionError:
        raise
    except Exception as exc:
        raise CVExtractionError(
            "The PDF appears damaged or could not be read. Export a fresh PDF and "
            "try again."
        ) from exc

    text = "\n".join(pages).strip()
    if not text:
        raise CVExtractionError(
            "No selectable text was found. The CV may be a scanned image; run OCR "
            "or upload a text-based PDF/DOCX."
        )
    return ExtractionResult(
        text=text[:MAX_TEXT_CHARACTERS],
        file_type="PDF",
        page_count=len(reader.pages),
        warnings=warnings,
    )


def _extract_docx(file_obj: BinaryIO) -> ExtractionResult:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - depends on deployment
        raise CVExtractionError(
            "DOCX analysis is temporarily unavailable because the document parser "
            "is not installed."
        ) from exc

    try:
        with ZipFile(file_obj) as archive:
            expanded_size = sum(item.file_size for item in archive.infolist())
            if expanded_size > 50 * 1024 * 1024:
                raise CVExtractionError(
                    "This DOCX expands to an unusually large document and cannot be "
                    "analysed safely."
                )
        file_obj.seek(0)
    except CVExtractionError:
        raise
    except (BadZipFile, OSError) as exc:
        raise CVExtractionError(
            "The DOCX appears damaged or is not a valid Word document."
        ) from exc

    try:
        document = Document(file_obj)
    except Exception as exc:
        raise CVExtractionError(
            "The DOCX appears damaged or could not be read. Save a fresh copy and "
            "try again."
        ) from exc

    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))

    text = "\n".join(blocks).strip()
    if not text:
        raise CVExtractionError("No readable text was found in this DOCX file.")
    warnings = []
    if len(document.tables) > 1:
        warnings.append(
            "Several tables were detected; some applicant tracking systems read "
            "table-based layouts out of order."
        )
    return ExtractionResult(
        text=text[:MAX_TEXT_CHARACTERS],
        file_type="DOCX",
        table_count=len(document.tables),
        warnings=warnings,
    )


def _normalise_lines(text: str) -> list[str]:
    return [
        re.sub(r"\s+", " ", line).strip()
        for line in text.replace("\x00", " ").splitlines()
        if line.strip()
    ]


def _has_section(lines: Iterable[str], aliases: Iterable[str]) -> bool:
    aliases = {alias.casefold() for alias in aliases}
    for line in lines:
        candidate = re.sub(r"[^a-zA-Z &]", "", line).strip().casefold()
        if candidate in aliases:
            return True
    return False


def _contains_term(text: str, term: str) -> bool:
    pattern = r"(?<![\w+#])" + re.escape(term) + r"(?![\w+#])"
    return bool(re.search(pattern, text, flags=re.IGNORECASE))


def _extract_skill_terms(text: str) -> set[str]:
    return {term for term in SKILL_TERMS if _contains_term(text, term)}


def _has_phone_number(text: str) -> bool:
    candidates = re.findall(
        r"(?<!\d)(?:\+\d{1,3}[\s.-]?)?(?:\d[\s().-]?){7,14}(?!\d)",
        text,
    )
    for candidate in candidates:
        digits = re.sub(r"\D", "", candidate)
        if not 7 <= len(digits) <= 15:
            continue
        # Do not mistake a common employment/education year range for a phone.
        if re.fullmatch(
            r"\s*(?:19|20)\d{2}\s*[-–—]\s*(?:19|20)\d{2}\s*",
            candidate,
        ):
            continue
        return True
    return False


def _top_context_keywords(text: str, limit: int = 12) -> list[str]:
    tokens = re.findall(r"[a-z][a-z+#.-]{2,}", text.casefold())
    counts = Counter(
        token.rstrip(".")
        for token in tokens
        if token.rstrip(".") not in STOP_WORDS and not token.isdigit()
    )
    return [token for token, _ in counts.most_common(limit)]


def _category(key: str, name: str, score: float, maximum: float, evidence: str) -> dict:
    percentage = round((score / maximum) * 100) if maximum else 0
    if percentage >= 85:
        status = "Excellent"
    elif percentage >= 70:
        status = "Strong"
    elif percentage >= 55:
        status = "Developing"
    else:
        status = "Needs attention"
    return {
        "key": key,
        "name": name,
        "score": round(score, 1),
        "maximum": maximum,
        "percentage": percentage,
        "status": status,
        "evidence": evidence,
    }


def _recommendation(
    priority: str,
    category: str,
    title: str,
    detail: str,
    action: str,
) -> dict:
    priority_order = {"High": 0, "Medium": 1, "Low": 2}
    return {
        "priority": priority,
        "priority_order": priority_order[priority],
        "category": category,
        "title": title,
        "detail": detail,
        "action": action,
    }


def analyse_cv_text(
    text: str,
    *,
    extraction: ExtractionResult | None = None,
    target_role: str = "",
    job_description: str = "",
    profile_skills: Iterable[str] | None = None,
) -> dict:
    """Return an evidence-backed scorecard and prioritised recommendations."""

    extraction = extraction or ExtractionResult(text=text, file_type="Text")
    clean_text = re.sub(r"[ \t]+", " ", text[:MAX_TEXT_CHARACTERS]).strip()
    lines = _normalise_lines(clean_text)
    lower_text = clean_text.casefold()
    words = re.findall(r"\b[\w+#.-]+\b", clean_text)
    word_count = len(words)

    sections = {
        key: _has_section(lines, aliases)
        for key, aliases in SECTION_ALIASES.items()
    }
    email_found = bool(
        re.search(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", clean_text, re.I)
    )
    phone_found = _has_phone_number(clean_text)
    linkedin_found = "linkedin.com/" in lower_text or "linkedin:" in lower_text
    portfolio_found = bool(
        re.search(
            r"(portfolio|behance\.net|dribbble\.com|github\.com/|https?://|www\.)",
            lower_text,
        )
    )

    bullet_lines = [
        line
        for line in lines
        if re.match(r"^\s*(?:[•●▪◦‣*-]|\d+[.)])\s*", line)
    ]
    achievement_lines = bullet_lines or [
        line for line in lines if any(line.casefold().startswith(verb) for verb in ACTION_VERBS)
    ]
    quantified_lines = [
        line
        for line in achievement_lines
        if re.search(
            r"(?:\b\d+(?:[.,]\d+)?\s*(?:%|percent|hours?|days?|weeks?|months?|"
            r"users?|clients?|customers?|projects?|people|members?|sales|revenue)\b|"
            r"(?:[$€£]|rs\.?|mur)\s*\d+)",
            line,
            re.I,
        )
    ]
    action_lines = [
        line
        for line in achievement_lines
        if re.match(
            r"^\s*(?:[•●▪◦‣*-]|\d+[.)])?\s*(?:"
            + "|".join(sorted(ACTION_VERBS))
            + r")\b",
            line,
            re.I,
        )
    ]
    achievement_count = len(achievement_lines)
    quantified_ratio = len(quantified_lines) / max(achievement_count, 1)
    action_ratio = len(action_lines) / max(achievement_count, 1)

    sentences = [
        sentence.strip()
        for sentence in re.split(r"(?<=[.!?])\s+", clean_text)
        if sentence.strip()
    ]
    average_sentence_words = (
        sum(len(sentence.split()) for sentence in sentences) / len(sentences)
        if sentences
        else 0
    )
    first_person_count = len(
        re.findall(r"\b(?:i|me|my|mine|we|our|ours)\b", lower_text)
    )
    generic_phrases = sorted(
        phrase for phrase in GENERIC_PHRASES if phrase in lower_text
    )
    unusual_character_ratio = (
        len(re.findall(r"[^\w\s.,;:!?%()+/#&@'’–—-]", clean_text))
        / max(len(clean_text), 1)
    )
    sensitive_details = sorted(
        label
        for label, pattern in {
            "date of birth": r"\b(?:date of birth|d\.?o\.?b\.?)\b",
            "marital status": r"\bmarital status\b",
            "national ID/passport": r"\b(?:national id|passport number)\b",
            "religion": r"\breligion\b",
        }.items()
        if re.search(pattern, lower_text)
    )

    completeness = 0.0
    completeness += 3 if email_found else 0
    completeness += 3 if phone_found else 0
    completeness += 2 if linkedin_found or portfolio_found else 0
    completeness += 2 if sections["summary"] else 0
    completeness += 4 if sections["experience"] else 0
    completeness += 3 if sections["education"] else 0
    completeness += 3 if sections["skills"] else 0

    essential_sections = sum(
        sections[key] for key in ("summary", "experience", "education", "skills")
    )
    structure = min(8.0, essential_sections * 2.0)
    structure += 3 if 350 <= word_count <= 950 else (1.5 if 200 <= word_count <= 1200 else 0)
    structure += 2 if achievement_count >= 4 else (1 if achievement_count else 0)
    structure += 2 if len(lines) >= 8 else 0

    impact = min(5.0, achievement_count * 0.75)
    impact += min(8.0, action_ratio * 10)
    impact += min(10.0, quantified_ratio * 20)
    impact += 2 if sections["projects"] or sections["certifications"] else 0

    readability = 0.0
    readability += 5 if 350 <= word_count <= 950 else (3 if 220 <= word_count <= 1100 else 1)
    readability += 4 if 0 < average_sentence_words <= 28 else (2 if average_sentence_words <= 35 else 0)
    readability += 3 if first_person_count <= 2 else (1 if first_person_count <= 6 else 0)
    readability += 3 if len(generic_phrases) <= 1 else (1 if len(generic_phrases) <= 3 else 0)

    ats = 5.0 if word_count >= 100 else 2.0
    ats += 3 if unusual_character_ratio < 0.01 else (1 if unusual_character_ratio < 0.03 else 0)
    ats += min(4.0, essential_sections)
    if extraction.table_count is None or extraction.table_count <= 1:
        ats += 3
    elif extraction.table_count == 2:
        ats += 1

    categories = [
        _category(
            "completeness",
            "Essential information",
            completeness,
            20,
            f"{sum((email_found, phone_found, linkedin_found or portfolio_found))}/3 "
            f"contact signals and {essential_sections}/4 core sections found.",
        ),
        _category(
            "structure",
            "Structure & focus",
            structure,
            15,
            f"{word_count} words, {essential_sections}/4 core sections and "
            f"{achievement_count} achievement-style lines.",
        ),
        _category(
            "impact",
            "Evidence & impact",
            impact,
            25,
            f"{len(action_lines)} action-led and {len(quantified_lines)} quantified "
            f"achievement lines detected.",
        ),
        _category(
            "readability",
            "Clarity & writing",
            readability,
            15,
            f"Average sentence length is {average_sentence_words:.0f} words; "
            f"{len(generic_phrases)} generic phrase(s) detected.",
        ),
        _category(
            "ats",
            "ATS compatibility",
            ats,
            15,
            f"Selectable text was extracted from {extraction.file_type}; "
            f"{essential_sections}/4 standard headings found.",
        ),
    ]

    cv_skills = _extract_skill_terms(clean_text)
    target_context = f"{target_role}\n{job_description}".strip()
    target_skills = _extract_skill_terms(target_context)
    context_keywords = _top_context_keywords(target_context) if job_description else []
    matched_skills = sorted(cv_skills & target_skills)
    missing_skills = sorted(target_skills - cv_skills)
    matched_keywords = [
        keyword for keyword in context_keywords if _contains_term(clean_text, keyword)
    ]

    if target_context:
        if target_skills:
            skill_alignment = len(matched_skills) / len(target_skills)
            alignment = min(7.0, skill_alignment * 7)
        else:
            alignment = 3.5
        keyword_alignment = len(matched_keywords) / max(len(context_keywords), 1)
        alignment += min(2.0, keyword_alignment * 3)
        role_words = [
            word
            for word in re.findall(r"[a-z][a-z+#.-]{2,}", target_role.casefold())
            if word not in STOP_WORDS
        ]
        if role_words and any(_contains_term(clean_text, word) for word in role_words):
            alignment += 1
        alignment = min(10.0, alignment)
        categories.append(
            _category(
                "alignment",
                "Role alignment",
                alignment,
                10,
                f"{len(matched_skills)}/{len(target_skills)} explicit target skills "
                f"and {len(matched_keywords)}/{len(context_keywords)} contextual "
                "keywords found.",
            )
        )

    total_score = sum(category["score"] for category in categories)
    total_maximum = sum(category["maximum"] for category in categories)
    overall_score = round((total_score / total_maximum) * 100)
    if overall_score >= 90:
        rating = "Excellent foundation"
    elif overall_score >= 80:
        rating = "Strong"
    elif overall_score >= 70:
        rating = "Competitive"
    elif overall_score >= 60:
        rating = "Developing"
    else:
        rating = "Needs attention"

    recommendations = []
    if not email_found or not phone_found:
        missing = []
        if not email_found:
            missing.append("professional email")
        if not phone_found:
            missing.append("phone number")
        recommendations.append(
            _recommendation(
                "High",
                "Contact",
                "Make it easy to contact you",
                f"The checker could not find: {', '.join(missing)}.",
                "Place contact details in plain text near your name, not only in a "
                "header, footer, image or icon.",
            )
        )
    missing_sections = [
        key.title()
        for key in ("summary", "experience", "education", "skills")
        if not sections[key]
    ]
    if missing_sections:
        recommendations.append(
            _recommendation(
                "High",
                "Structure",
                "Add recognisable section headings",
                f"Missing or unrecognised: {', '.join(missing_sections)}.",
                "Use simple headings such as Professional Summary, Experience, "
                "Education and Skills so recruiters and ATS tools can scan quickly.",
            )
        )
    if achievement_count < 4:
        recommendations.append(
            _recommendation(
                "High",
                "Impact",
                "Turn responsibilities into achievement bullets",
                f"Only {achievement_count} achievement-style line(s) were detected.",
                "Add 3–5 concise bullets per recent role using: action + task + "
                "result. Example: “Automated weekly reporting, saving 4 hours.”",
            )
        )
    elif quantified_ratio < 0.25:
        recommendations.append(
            _recommendation(
                "High",
                "Impact",
                "Prove outcomes with credible numbers",
                f"{len(quantified_lines)} of {achievement_count} achievement lines "
                "include a measurable result.",
                "Where truthful, add scale, time saved, quality, revenue, users, "
                "volume or percentage change. Never invent metrics.",
            )
        )
    if action_ratio < 0.5 and achievement_count:
        recommendations.append(
            _recommendation(
                "Medium",
                "Writing",
                "Lead bullets with specific action verbs",
                f"{len(action_lines)} of {achievement_count} achievement lines start "
                "with a strong action verb.",
                "Replace openers such as “Responsible for” with verbs such as "
                "Built, Analysed, Improved, Delivered or Led.",
            )
        )
    if word_count < 250:
        recommendations.append(
            _recommendation(
                "High",
                "Depth",
                "Add enough evidence to assess your fit",
                f"The extracted CV contains only {word_count} words.",
                "Add relevant projects, achievements, tools and outcomes. Keep every "
                "line useful rather than adding filler.",
            )
        )
    elif word_count > 1100:
        recommendations.append(
            _recommendation(
                "Medium",
                "Focus",
                "Tighten the CV",
                f"The extracted CV contains {word_count} words.",
                "Prioritise recent, role-relevant evidence; remove repetition and "
                "compress older experience.",
            )
        )
    if len(generic_phrases) > 1:
        recommendations.append(
            _recommendation(
                "Medium",
                "Credibility",
                "Replace generic claims with proof",
                f"Generic phrases detected: {', '.join(generic_phrases[:4])}.",
                "Show the behaviour through a concrete example or result instead of "
                "describing yourself with unsupported adjectives.",
            )
        )
    if first_person_count > 6:
        recommendations.append(
            _recommendation(
                "Low",
                "Writing",
                "Use a concise CV voice",
                f"{first_person_count} first-person pronouns were detected.",
                "Remove repeated “I”, “my” and “we” from bullets; start directly with "
                "the action.",
            )
        )
    if extraction.table_count and extraction.table_count > 1:
        recommendations.append(
            _recommendation(
                "Medium",
                "ATS",
                "Simplify the document layout",
                f"{extraction.table_count} tables were detected in the DOCX.",
                "Use a single-column layout with normal paragraphs where possible, "
                "then verify the reading order by copying the text into a plain editor.",
            )
        )
    if unusual_character_ratio >= 0.03:
        recommendations.append(
            _recommendation(
                "Medium",
                "ATS",
                "Check text extraction and icon usage",
                "A high proportion of unusual symbols was detected.",
                "Replace icon-only labels with words and confirm that copied text "
                "appears in the correct order.",
            )
        )
    if sensitive_details:
        recommendations.append(
            _recommendation(
                "Medium",
                "Privacy",
                "Remove unnecessary sensitive information",
                f"Potentially sensitive details detected: {', '.join(sensitive_details)}.",
                "Unless a specific application legally requires them, omit personal "
                "details that do not demonstrate job fit.",
            )
        )
    if target_context and missing_skills:
        recommendations.append(
            _recommendation(
                "High",
                "Role alignment",
                f"Address the largest gaps for {target_role or 'the target role'}",
                "Target skills not found in the CV: " + ", ".join(missing_skills[:8]) + ".",
                "Add only skills you genuinely have and support important ones with "
                "evidence in Experience or Projects—not just a keyword list.",
            )
        )
    elif target_context and not target_skills:
        recommendations.append(
            _recommendation(
                "Low",
                "Role alignment",
                "Use a fuller job description",
                "The target text did not contain enough recognisable technical or "
                "professional skills for a precise comparison.",
                "Paste the responsibilities and requirements from the job advert for "
                "a more useful role-alignment check.",
            )
        )

    if profile_skills:
        profile_skill_names = {
            str(skill).strip().casefold()
            for skill in profile_skills
            if str(skill).strip()
        }
        missing_profile_skills = sorted(
            skill for skill in profile_skill_names if not _contains_term(clean_text, skill)
        )
        if missing_profile_skills:
            recommendations.append(
                _recommendation(
                    "Low",
                    "Profile consistency",
                    "Reconcile your profile and CV skills",
                    "Profile skills not found in this CV: "
                    + ", ".join(missing_profile_skills[:6])
                    + ".",
                    "If these skills are current and relevant, demonstrate them in a "
                    "project or experience bullet; otherwise update your profile.",
                )
            )

    if not recommendations:
        recommendations.append(
            _recommendation(
                "Low",
                "Quality assurance",
                "Complete a final human review",
                "No major structural issues were detected by the automated checks.",
                "Ask a trusted reviewer in your field to verify accuracy, tone, "
                "formatting and role relevance before submitting.",
            )
        )

    recommendations.sort(
        key=lambda item: (item["priority_order"], item["category"], item["title"])
    )
    recommendations = recommendations[:8]

    strengths = []
    if email_found and phone_found:
        strengths.append("Recruiters can find both an email address and phone number.")
    if essential_sections == 4:
        strengths.append("All four standard CV sections are clearly labelled.")
    if achievement_count >= 6:
        strengths.append(
            f"The CV uses {achievement_count} achievement-style lines for scanability."
        )
    if quantified_ratio >= 0.3:
        strengths.append("A meaningful share of achievements is backed by numbers.")
    if 350 <= word_count <= 950:
        strengths.append("The content length is within a focused professional range.")
    if target_context and matched_skills:
        strengths.append(
            "Explicit target-skill matches include " + ", ".join(matched_skills[:6]) + "."
        )
    if not strengths:
        strengths.append(
            "The document contains selectable text, so the checker could inspect its content."
        )

    if word_count >= 450 and essential_sections >= 3:
        confidence = "High"
    elif word_count >= 180:
        confidence = "Moderate"
    else:
        confidence = "Low"

    checks = [
        {"label": "Professional email", "passed": email_found},
        {"label": "Phone number", "passed": phone_found},
        {"label": "Experience section", "passed": sections["experience"]},
        {"label": "Education section", "passed": sections["education"]},
        {"label": "Skills section", "passed": sections["skills"]},
        {"label": "Four or more achievement lines", "passed": achievement_count >= 4},
        {"label": "Measurable outcomes", "passed": bool(quantified_lines)},
        {
            "label": "Focused content length",
            "passed": 250 <= word_count <= 1100,
        },
    ]

    return {
        "version": ANALYZER_VERSION,
        "overall_score": overall_score,
        "rating": rating,
        "confidence": confidence,
        "categories": categories,
        "recommendations": recommendations,
        "strengths": strengths[:6],
        "checks": checks,
        "metrics": {
            "word_count": word_count,
            "achievement_count": achievement_count,
            "quantified_count": len(quantified_lines),
            "target_skill_count": len(target_skills),
            "matched_skill_count": len(matched_skills),
        },
        "target": {
            "role": target_role.strip(),
            "matched_skills": matched_skills,
            "missing_skills": missing_skills,
        },
        "extraction": {
            "file_type": extraction.file_type,
            "page_count": extraction.page_count,
            "table_count": extraction.table_count,
            "warnings": extraction.warnings,
        },
        "disclaimer": (
            "This is an explainable automated review, not a hiring decision or a "
            "guarantee of ATS performance. Scores are directional; verify advice "
            "against the specific employer and have a person review the final CV."
        ),
    }
