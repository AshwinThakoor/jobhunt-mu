"""Server-rendered request handlers for the JobHunt MU web application."""
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, JsonResponse, FileResponse, Http404
from django.urls import reverse
from django.utils.http import url_has_allowed_host_and_scheme
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib import messages
from django.db import transaction
from django.db.models import Q, Count, Avg
from django.core.paginator import Paginator
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from django.conf import settings
from decimal import Decimal
import json
import stripe
from datetime import datetime, timedelta
import random
import os

from .models import (
    UserProfile, CV, Company, Internship, Application, 
    Payment, Recommendation, SavedJob, CareerDocument, SearchHistory, StudentRecord, ModuleMark
)
from .forms import (
    UserRegistrationForm, UserProfileForm, CVUploadForm, 
    InternshipSearchForm, ApplicationForm, CompanyForm, 
    InternshipForm, PaymentForm, CVAnalysisForm
)
from .services.cv_analyzer import (
    CVExtractionError,
    analyse_cv_text,
    extract_cv,
)
from .services.job_matcher import extract_known_skills, match_resume_to_job
from .services.application_studio import build_application_pack

# Initialize Stripe
stripe.api_key = settings.STRIPE_SECRET_KEY
StripeError = getattr(
    stripe,
    "StripeError",
    getattr(getattr(stripe, "error", None), "StripeError", Exception),
)
StripeSignatureError = getattr(
    stripe,
    "SignatureVerificationError",
    getattr(getattr(stripe, "error", None), "SignatureVerificationError", Exception),
)

SOURCE_DIRECTORY = [
    {
        'name': 'MyJob.mu',
        'kind': 'Mauritius',
        'description': 'Detailed local vacancies and internships from Mauritius employers.',
        'url': 'https://www.myjob.mu/jobs',
        'access': 'automatic',
    },
    {
        'name': 'Jobs.mu',
        'kind': 'Mauritius',
        'description': 'Local jobs across industries, experience levels and contract types.',
        'url': 'https://www.jobs.mu/jobs/',
        'access': 'automatic',
    },
    {
        'name': 'Mauritius Jobs',
        'kind': 'Government',
        'description': 'Vacancies published through the official national employment portal.',
        'url': 'https://mauritiusjobs.govmu.org/jobsearch',
        'access': 'automatic',
    },
    {
        'name': 'Remotive',
        'kind': 'Remote',
        'description': 'Remote roles from global employers, always linked back to Remotive.',
        'url': 'https://remotive.com/remote-jobs',
        'access': 'automatic',
    },
    {
        'name': 'Freelancer.com',
        'kind': 'Freelance',
        'description': 'Global freelance projects through the official Freelancer API.',
        'url': 'https://www.freelancer.com/jobs/',
        'access': 'api_key',
    },
    {
        'name': 'Jooble',
        'kind': 'Aggregator',
        'description': 'Additional Mauritius vacancies through an approved Jooble API key.',
        'url': 'https://jooble.org/api/about',
        'access': 'api_key',
    },
    {
        'name': 'LinkedIn',
        'kind': 'Professional',
        'description': 'Search Mauritius roles directly on LinkedIn without copying restricted data.',
        'url': 'https://www.linkedin.com/jobs/search/?location=Mauritius',
        'access': 'external',
    },
    {
        'name': 'Upwork',
        'kind': 'Freelance',
        'description': 'Browse freelance projects directly; API importing requires approved access.',
        'url': 'https://www.upwork.com/nx/search/jobs/',
        'access': 'external',
    },
    {
        'name': 'Toptal',
        'kind': 'Freelance',
        'description': 'Apply to the talent network to access matched premium freelance projects.',
        'url': 'https://www.toptal.com/talent/apply',
        'access': 'external',
    },
    {
        'name': 'CareerHub',
        'kind': 'Mauritius',
        'description': 'Discover additional local employers and roles directly on CareerHub.',
        'url': 'https://www.careerhub.mu/',
        'access': 'external',
    },
]

# Authentication Views
def login_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')
    
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        
        if user is not None:
            login(request, user)
            messages.success(request, f'Welcome back, {user.first_name}!')
            next_url = request.POST.get('next') or request.GET.get('next')
            if next_url and url_has_allowed_host_and_scheme(
                next_url,
                allowed_hosts={request.get_host()},
                require_https=request.is_secure(),
            ):
                return redirect(next_url)
            return redirect('dashboard')
        else:
            messages.error(request, 'Invalid username or password')
    
    return render(request, 'myapp/login.html', {'next': request.GET.get('next', '')})

def signup_view(request):
    if request.user.is_authenticated:
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = UserRegistrationForm(request.POST)
        if form.is_valid():
            user = form.save()
            messages.success(request, 'Account created successfully! You can now log in.')
            return redirect('login')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = UserRegistrationForm()
    
    return render(request, 'myapp/signup.html', {'form': form})

def logout_view(request):
    logout(request)
    messages.info(request, 'You have been logged out.')
    return redirect('login')

# Dashboard and Main Views
@login_required
def dashboard(request):
    user_profile, _ = UserProfile.objects.get_or_create(user=request.user)

    if user_profile.user_type == 'student':
        applications_qs = Application.objects.filter(user=request.user).select_related('internship', 'internship__company')
        cvs = CV.objects.filter(user=request.user).order_by('-is_primary', '-uploaded_at')
        recommendations_qs = Recommendation.objects.filter(user=request.user).exclude(status='dismissed').select_related('internship', 'internship__company').order_by('-score')
        saved_count = SavedJob.objects.filter(user=request.user).count()

        profile_checks = {
            'name': bool(request.user.first_name and request.user.last_name),
            'contact': bool(request.user.email and user_profile.phone),
            'location': bool(user_profile.address),
            'summary': bool(user_profile.bio and len(user_profile.bio.strip()) >= 40),
            'skills': bool(user_profile.skills),
            'education': bool(user_profile.education),
            'experience': bool(user_profile.experience),
            'resume': cvs.exists(),
        }
        profile_completion = round(sum(profile_checks.values()) / len(profile_checks) * 100)
        pipeline = {
            key: applications_qs.filter(status=key).count()
            for key, _ in Application.STATUS_CHOICES
        }
        context = {
            'user_profile': user_profile,
            'applications': applications_qs.order_by('-updated_at')[:5],
            'cvs': cvs[:4],
            'recommendations': recommendations_qs[:6] if user_profile.is_premium else recommendations_qs[:3],
            'total_applications': applications_qs.count(),
            'accepted_applications': applications_qs.filter(status='accepted').count(),
            'saved_count': saved_count,
            'profile_completion': profile_completion,
            'profile_checks': profile_checks,
            'pipeline': pipeline,
            'top_match': recommendations_qs.first(),
            'new_matches': recommendations_qs.filter(created_at__gte=timezone.now()-timedelta(days=7)).count(),
            'is_premium': user_profile.is_premium,
        }
    else:
        company = Company.objects.filter(name__icontains=request.user.first_name).first()
        internships = Internship.objects.filter(company=company).order_by('-created_at')[:5] if company else []
        context = {
            'user_profile': user_profile,
            'internships': internships,
            'total_applications': Application.objects.filter(internship__company=company).count() if company else 0,
        }
    return render(request, 'myapp/dashboard.html', context)

@login_required
def profile_view(request):
    user_profile, created = UserProfile.objects.get_or_create(user=request.user)
    
    if request.method == 'POST':
        form = UserProfileForm(request.POST, request.FILES, instance=user_profile)
        if form.is_valid():
            form.save()
            messages.success(request, 'Profile updated successfully!')
            return redirect('profile')
    else:
        form = UserProfileForm(instance=user_profile)
    
    context = {
        'form': form,
        'user_profile': user_profile,
    }
    return render(request, 'myapp/profile.html', context)

# Internship Views
def internship_list(request):
    search_form = InternshipSearchForm(request.GET)
    base_internships = Internship.objects.filter(status='active').select_related('company')
    internships = base_internships
    
    if search_form.is_valid():
        query = search_form.cleaned_data.get('query')
        location = search_form.cleaned_data.get('location')
        duration = search_form.cleaned_data.get('duration')
        stipend_min = search_form.cleaned_data.get('stipend_min')
        skills = search_form.cleaned_data.get('skills')
        opportunity_type = search_form.cleaned_data.get('opportunity_type')
        source = search_form.cleaned_data.get('source')
        
        if query:
            internships = internships.filter(
                Q(title__icontains=query) | 
                Q(description__icontains=query) |
                Q(company__name__icontains=query)
            )
        
        if location:
            internships = internships.filter(location__icontains=location)
        
        if duration:
            internships = internships.filter(duration__icontains=duration)
        
        if stipend_min:
            # This is a simplified filter - in real app you'd parse stipend amounts
            internships = internships.filter(stipend__isnull=False)
        
        if skills:
            skill_list = [skill.strip() for skill in skills.split(',')]
            for skill in skill_list:
                internships = internships.filter(skills_required__contains=[skill])

        if opportunity_type:
            internships = internships.filter(opportunity_type=opportunity_type)

        if source:
            internships = internships.filter(source_name=source)
        
        # Save search history for logged-in users
        if request.user.is_authenticated and any([
            query, location, duration, stipend_min, skills, opportunity_type, source
        ]):
            serializable_filters = {
                key: float(value) if isinstance(value, Decimal) else value
                for key, value in search_form.cleaned_data.items()
            }
            SearchHistory.objects.create(
                user=request.user,
                query=query or '',
                filters=serializable_filters
            )
    
    total_count = internships.count()
    company_count = internships.values('company_id').distinct().count()
    active_source_count = base_internships.values('source_name').distinct().count()
    source_counts = {
        row['source_name']: row['count']
        for row in base_internships.values('source_name').annotate(count=Count('id'))
    }
    source_directory = [
        {**item, 'count': source_counts.get(item['name'], 0)}
        for item in SOURCE_DIRECTORY
    ]
    category_cards = [
        {
            'key': key,
            'label': label,
            'count': base_internships.filter(opportunity_type=key).count(),
        }
        for key, label in Internship.OPPORTUNITY_TYPE_CHOICES
    ]
    query_params = request.GET.copy()
    query_params.pop('page', None)

    # Pagination
    paginator = Paginator(internships, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'internships': page_obj,
        'search_form': search_form,
        'total_count': total_count,
        'company_count': company_count,
        'active_source_count': active_source_count,
        'querystring': query_params.urlencode(),
        'source_directory': source_directory,
        'category_cards': category_cards,
    }
    return render(request, 'myapp/internship_list.html', context)

def internship_detail(request, pk):
    internship = get_object_or_404(
        Internship.objects.select_related('company'),
        pk=pk,
    )
    user_has_applied = False
    user_is_premium = False
    
    if request.user.is_authenticated:
        user_has_applied = Application.objects.filter(
            user=request.user, 
            internship=internship
        ).exists()
        profile, _ = UserProfile.objects.get_or_create(user=request.user)
        user_is_premium = profile.is_premium

    can_view_details = not internship.is_premium or user_is_premium
    related_internships = Internship.objects.filter(
        status='active',
        company=internship.company,
    ).exclude(pk=internship.pk)[:3]
    
    context = {
        'internship': internship,
        'user_has_applied': user_has_applied,
        'can_view_details': can_view_details,
        'related_internships': related_internships,
    }
    return render(request, 'myapp/internship_detail.html', context)

@login_required
def apply_internship(request, pk):
    internship = get_object_or_404(Internship, pk=pk)
    
    # Check if user already applied
    if Application.objects.filter(user=request.user, internship=internship).exists():
        messages.warning(request, 'You have already applied for this internship.')
        return redirect('internship_detail', pk=pk)
    
    # Check if internship is still open
    if not internship.is_open:
        messages.error(request, 'This internship is no longer accepting applications.')
        return redirect('internship_detail', pk=pk)
    
    if request.method == 'POST':
        form = ApplicationForm(request.user, request.POST)
        if form.is_valid():
            application = form.save(commit=False)
            application.user = request.user
            application.internship = internship
            application.save()
            
            messages.success(request, 'Application submitted successfully!')
            return redirect('application_detail', pk=application.pk)
    else:
        form = ApplicationForm(request.user)
    
    context = {
        'form': form,
        'internship': internship,
    }
    return render(request, 'myapp/apply_internship.html', context)

# CV Management Views
@login_required
def cv_list(request):
    cvs = CV.objects.filter(user=request.user).order_by('-is_primary', '-uploaded_at')
    if request.method == 'POST':
        form = CVUploadForm(request.POST, request.FILES)
        if form.is_valid():
            cv = form.save(commit=False)
            cv.user = request.user
            cv.is_primary = not CV.objects.filter(user=request.user).exists()
            cv.save()
            try:
                cv.file.open('rb')
                extraction = extract_cv(cv.file, cv.file.name)
                skills = extract_known_skills(extraction.text)
                cv.extracted_text = extraction.text
                cv.extracted_skills = skills
                cv.extraction_warnings = extraction.warnings
                cv.parsed_at = timezone.now()
                cv.validated = bool(extraction.text.strip())
                cv.validation_notes = 'Resume text extracted successfully.' if cv.validated else 'No readable text found.'
                cv.save(update_fields=['extracted_text','extracted_skills','extraction_warnings','parsed_at','validated','validation_notes','is_primary'])
                profile, _ = UserProfile.objects.get_or_create(user=request.user)
                profile.skills = sorted(set((profile.skills or []) + skills))
                profile.save(update_fields=['skills','updated_at'])
                messages.success(request, f'CV uploaded and analysed. {len(skills)} skills detected.')
            except CVExtractionError as exc:
                cv.validation_notes = str(exc)
                cv.save(update_fields=['validation_notes'])
                messages.warning(request, f'CV uploaded, but automatic extraction needs attention: {exc}')
            finally:
                try: cv.file.close()
                except Exception: pass
            return redirect('cv_list')
    else:
        form = CVUploadForm()
    return render(request, 'myapp/cv_list.html', {'cvs': cvs, 'form': form})

@login_required
def cv_detail(request, pk):
    cv = get_object_or_404(CV, pk=pk, user=request.user)
    return render(request, 'myapp/cv_detail.html', {'cv': cv})


@login_required
def cv_analyze(request, pk):
    cv = get_object_or_404(CV, pk=pk, user=request.user)
    form = CVAnalysisForm(request.POST or None)
    report = None
    analysis_error = None

    if request.method == 'POST' and form.is_valid():
        try:
            if cv.file.size > 10 * 1024 * 1024:
                raise CVExtractionError(
                    "This file is too large to analyse safely. Upload a CV under 10 MB."
                )
            cv.file.open('rb')
            extraction = extract_cv(cv.file, cv.file.name)
            user_profile, _ = UserProfile.objects.get_or_create(user=request.user)
            report = analyse_cv_text(
                extraction.text,
                extraction=extraction,
                target_role=form.cleaned_data['target_role'],
                job_description=form.cleaned_data['job_description'],
                profile_skills=user_profile.skills or [],
            )
        except (CVExtractionError, OSError, ValueError) as exc:
            analysis_error = str(exc)
        finally:
            try:
                cv.file.close()
            except (AttributeError, OSError):
                pass

    return render(
        request,
        'myapp/cv_analyze.html',
        {
            'cv': cv,
            'form': form,
            'report': report,
            'analysis_error': analysis_error,
        },
    )


@login_required
def cv_delete(request, pk):
    cv = get_object_or_404(CV, pk=pk, user=request.user)
    if request.method == 'POST':
        cv.delete()
        messages.success(request, 'CV deleted successfully!')
        return redirect('cv_list')
    return render(request, 'myapp/cv_confirm_delete.html', {'cv': cv})

@login_required
def set_primary_cv(request, pk):
    cv = get_object_or_404(CV, pk=pk, user=request.user)
    cv.is_primary = True
    cv.save()
    messages.success(request, f'"{cv.title}" is now your primary CV.')
    return redirect('cv_list')

# Application Tracking Views
@login_required
def application_list(request):
    applications = Application.objects.filter(user=request.user).order_by('-applied_at')
    
    # Filter by status
    status_filter = request.GET.get('status')
    if status_filter:
        applications = applications.filter(status=status_filter)
    
    context = {
        'applications': applications,
        'status_choices': Application.STATUS_CHOICES,
        'current_status': status_filter,
    }
    return render(request, 'myapp/application_list.html', context)

@login_required
def application_detail(request, pk):
    application = get_object_or_404(Application, pk=pk, user=request.user)
    return render(request, 'myapp/application_detail.html', {'application': application})

@login_required
def withdraw_application(request, pk):
    application = get_object_or_404(Application, pk=pk, user=request.user)
    if request.method == 'POST':
        application.status = 'withdrawn'
        application.save()
        messages.success(request, 'Application withdrawn successfully!')
        return redirect('application_list')
    return render(request, 'myapp/withdraw_application.html', {'application': application})

# Payment Views
@login_required
def payment_list(request):
    payments = Payment.objects.filter(user=request.user).order_by('-created_at')
    return render(request, 'myapp/payment_list.html', {'payments': payments})


def _stripe_value(obj, key, default=None):
    if isinstance(obj, dict):
        return obj.get(key, default)
    return getattr(obj, key, default)


@transaction.atomic
def _complete_premium_payment(payment, payment_intent_id=None):
    if payment.status != 'completed':
        payment.status = 'completed'
        payment.completed_at = timezone.now()
        if payment_intent_id:
            payment.stripe_payment_intent_id = str(payment_intent_id)
        payment.save(
            update_fields=['status', 'completed_at', 'stripe_payment_intent_id']
        )

    profile, _ = UserProfile.objects.get_or_create(user=payment.user)
    if not profile.is_premium:
        profile.is_premium = True
        profile.save(update_fields=['is_premium'])


@login_required
def create_payment(request):
    premium_cents = settings.PREMIUM_PRICE_CENTS
    premium_amount = Decimal(premium_cents) / Decimal('100')
    currency = settings.PREMIUM_CURRENCY.lower()

    profile, _ = UserProfile.objects.get_or_create(user=request.user)
    if profile.is_premium:
        messages.info(request, 'Your account already has Premium access.')
        return redirect('dashboard')

    if request.method == 'POST':
        if not settings.STRIPE_SECRET_KEY:
            messages.error(
                request,
                'Payments are not configured yet. Add your Stripe keys and try again.'
            )
            return redirect('create_payment')

        payment = Payment.objects.create(
            user=request.user,
            payment_type='premium_subscription',
            amount=premium_amount,
            currency=currency.upper(),
            description='JobHunt MU Premium access',
        )
        metadata = {
            'user_id': str(request.user.pk),
            'payment_id': str(payment.pk),
            'payment_type': 'premium_subscription',
        }
        checkout_args = {
            'mode': 'payment',
            'payment_method_types': ['card'],
            'client_reference_id': str(request.user.pk),
            'line_items': [{
                'price_data': {
                    'currency': currency,
                    'product_data': {
                        'name': 'JobHunt MU Premium',
                        'description': 'Premium recommendations, application insights and CV analytics',
                    },
                    'unit_amount': premium_cents,
                },
                'quantity': 1,
            }],
            'success_url': (
                request.build_absolute_uri(reverse('payment_success'))
                + '?session_id={CHECKOUT_SESSION_ID}'
            ),
            'cancel_url': request.build_absolute_uri(reverse('create_payment')),
            'metadata': metadata,
            'payment_intent_data': {'metadata': metadata},
        }
        if request.user.email:
            checkout_args['customer_email'] = request.user.email

        try:
            checkout_session = stripe.checkout.Session.create(**checkout_args)
        except StripeError:
            payment.status = 'failed'
            payment.save(update_fields=['status'])
            messages.error(
                request,
                'Stripe could not start the checkout. Please try again in a moment.'
            )
            return redirect('create_payment')

        payment.stripe_checkout_session_id = _stripe_value(checkout_session, 'id')
        payment.save(update_fields=['stripe_checkout_session_id'])
        return redirect(_stripe_value(checkout_session, 'url'))

    return render(request, 'myapp/create_payment.html', {
        'premium_amount': premium_amount,
        'premium_currency': currency.upper(),
        'stripe_ready': bool(settings.STRIPE_SECRET_KEY),
    })


@login_required
def payment_success(request):
    session_id = request.GET.get('session_id', '').strip()
    if not session_id:
        messages.warning(request, 'We could not verify that checkout session.')
        return redirect('payment_list')

    try:
        checkout_session = stripe.checkout.Session.retrieve(session_id)
    except StripeError:
        messages.error(
            request,
            'We could not verify your payment with Stripe. Please contact support if you were charged.'
        )
        return redirect('payment_list')

    if (
        str(_stripe_value(checkout_session, 'client_reference_id')) != str(request.user.pk)
        or _stripe_value(checkout_session, 'payment_status') != 'paid'
    ):
        messages.warning(request, 'Your payment has not been confirmed yet.')
        return redirect('payment_list')

    payment = get_object_or_404(
        Payment,
        user=request.user,
        stripe_checkout_session_id=session_id,
        payment_type='premium_subscription',
    )
    _complete_premium_payment(
        payment,
        _stripe_value(checkout_session, 'payment_intent'),
    )
    return render(request, 'myapp/payment_success.html', {'payment': payment})


@csrf_exempt
@require_POST
def stripe_webhook(request):
    if not settings.STRIPE_WEBHOOK_SECRET:
        return HttpResponse('Stripe webhook secret is not configured.', status=503)

    sig_header = request.META.get('HTTP_STRIPE_SIGNATURE')
    if not sig_header:
        return HttpResponse(status=400)
    try:
        event = stripe.Webhook.construct_event(
            request.body,
            sig_header,
            settings.STRIPE_WEBHOOK_SECRET,
        )
    except (ValueError, StripeSignatureError):
        return HttpResponse(status=400)

    event_type = _stripe_value(event, 'type')
    event_data = _stripe_value(_stripe_value(event, 'data', {}), 'object', {})
    metadata = _stripe_value(event_data, 'metadata', {}) or {}
    payment_id = _stripe_value(metadata, 'payment_id')

    if (
        event_type in {'checkout.session.completed', 'checkout.session.async_payment_succeeded'}
        and _stripe_value(event_data, 'payment_status') == 'paid'
        and payment_id
    ):
        try:
            payment = Payment.objects.select_related('user').get(
                pk=payment_id,
                stripe_checkout_session_id=_stripe_value(event_data, 'id'),
            )
        except Payment.DoesNotExist:
            return HttpResponse(status=200)
        _complete_premium_payment(
            payment,
            _stripe_value(event_data, 'payment_intent'),
        )

    elif event_type == 'payment_intent.payment_failed' and payment_id:
        Payment.objects.filter(pk=payment_id, status='pending').update(status='failed')

    return HttpResponse(status=200)

# Recommendation Engine
@login_required
def recommendations(request):
    profile, _ = UserProfile.objects.get_or_create(user=request.user)
    primary_cv = CV.objects.filter(user=request.user, is_primary=True).first() or CV.objects.filter(user=request.user).order_by('-uploaded_at').first()
    if not primary_cv:
        messages.info(request, 'Upload a CV first so JobHunt MU can explain your job matches.')
        return redirect('cv_list')

    resume_text = primary_cv.extracted_text
    if not resume_text:
        try:
            primary_cv.file.open('rb')
            extraction = extract_cv(primary_cv.file, primary_cv.file.name)
            resume_text = extraction.text
            primary_cv.extracted_text = resume_text
            primary_cv.extracted_skills = extract_known_skills(resume_text)
            primary_cv.parsed_at = timezone.now()
            primary_cv.save(update_fields=['extracted_text','extracted_skills','parsed_at'])
        except CVExtractionError as exc:
            messages.error(request, str(exc))
            return redirect('cv_analyze', pk=primary_cv.pk)
        finally:
            try: primary_cv.file.close()
            except Exception: pass

    active_jobs = Internship.objects.filter(status='active').select_related('company')
    for job in active_jobs:
        result = match_resume_to_job(resume_text=resume_text, profile_skills=profile.skills or primary_cv.extracted_skills, job=job)
        recommendation_defaults = {
            'score': result['score'],
            'reason': result['reason'],
            'matched_skills': result['matched_skills'],
            'missing_skills': result['missing_skills'],
            'score_breakdown': {**result['score_breakdown'], 'related_skills': result.get('related_skills', [])},
            'match_confidence': result.get('match_confidence', 'medium'),
        }
        rec, _ = Recommendation.objects.get_or_create(user=request.user, internship=job, defaults=recommendation_defaults)
        if rec.status != 'dismissed':
            rec.score = recommendation_defaults['score']; rec.reason = recommendation_defaults['reason']
            rec.matched_skills = recommendation_defaults['matched_skills']; rec.missing_skills = recommendation_defaults['missing_skills']
            rec.score_breakdown = recommendation_defaults['score_breakdown']
            rec.match_confidence = recommendation_defaults['match_confidence']
            rec.save(update_fields=['score','reason','matched_skills','missing_skills','score_breakdown','match_confidence','updated_at'])

    status = request.GET.get('status', 'active')
    qs = Recommendation.objects.filter(user=request.user).select_related('internship','internship__company')
    if status == 'saved':
        qs = qs.filter(status='saved')
    elif status == 'all':
        qs = qs.exclude(status='dismissed')
    else:
        qs = qs.filter(status='active')
    limit = 30 if profile.is_premium else 3
    return render(request, 'myapp/recommendations.html', {
        'recommendations': qs.order_by('-score')[:limit], 'primary_cv': primary_cv,
        'status_filter': status, 'is_limited': not profile.is_premium,
        'is_premium': profile.is_premium,
        'visible_limit': limit,
    })


@login_required
@require_POST
def recommendation_action(request, pk):
    recommendation = get_object_or_404(Recommendation, pk=pk, user=request.user)
    action = request.POST.get('action')
    if action == 'save':
        recommendation.status = 'saved'
        SavedJob.objects.get_or_create(user=request.user, internship=recommendation.internship)
        message = 'Job saved to your career dashboard.'
    elif action == 'dismiss':
        recommendation.status = 'dismissed'
        SavedJob.objects.filter(user=request.user, internship=recommendation.internship).delete()
        message = 'Recommendation hidden. This feedback will help refine your feed.'
    elif action == 'restore':
        recommendation.status = 'active'; message = 'Recommendation restored.'
    else:
        return JsonResponse({'ok': False, 'error': 'Unknown action'}, status=400)
    recommendation.save(update_fields=['status','updated_at'])
    messages.success(request, message)
    return redirect(request.POST.get('next') or 'recommendations')


@login_required
def saved_jobs(request):
    items = SavedJob.objects.filter(user=request.user).select_related('internship','internship__company')
    return render(request, 'myapp/saved_jobs.html', {'saved_jobs': items})

# Roadmap Views
@login_required
def roadmap(request):
    user_profile, created = UserProfile.objects.get_or_create(user=request.user)
    applications = Application.objects.filter(user=request.user)
    
    # Calculate profile completion percentage
    profile_fields = [
        user_profile.phone, user_profile.address, user_profile.bio,
        user_profile.skills, user_profile.education, user_profile.experience
    ]
    completed_fields = sum(1 for field in profile_fields if field)
    profile_completion = (completed_fields / len(profile_fields)) * 100
    
    # CV status
    cvs = CV.objects.filter(user=request.user)
    has_primary_cv = cvs.filter(is_primary=True).exists()
    
    # Application statistics
    total_applications = applications.count()
    accepted_applications = applications.filter(status='accepted').count()
    success_rate = (accepted_applications / total_applications * 100) if total_applications > 0 else 0
    
    # Recommendations for improvement
    recommendations = []
    
    if profile_completion < 80:
        recommendations.append("Complete your profile to increase your chances")
    
    if not has_primary_cv:
        recommendations.append("Upload and set a primary CV")
    
    if total_applications < 5:
        recommendations.append("Apply to more internships to increase your chances")
    
    if success_rate < 20 and total_applications > 0:
        recommendations.append("Consider improving your CV and cover letter")
    
    context = {
        'profile_completion': profile_completion,
        'has_primary_cv': has_primary_cv,
        'total_applications': total_applications,
        'accepted_applications': accepted_applications,
        'success_rate': success_rate,
        'recommendations': recommendations,
    }
    return render(request, 'myapp/roadmap.html', context)

# Admin/Employer Views
@login_required
def employer_dashboard(request):
    user_profile, created = UserProfile.objects.get_or_create(user=request.user)
    if user_profile.user_type != 'employer':
        messages.error(request, 'Access denied. Employer account required.')
        return redirect('dashboard')
    
    # Get company (simplified - in real app you'd have proper company association)
    company = Company.objects.filter(name__icontains=request.user.first_name).first()
    
    if not company:
        messages.warning(request, 'Please create a company profile first.')
        return redirect('create_company')
    
    internships = Internship.objects.filter(company=company)
    total_applications = Application.objects.filter(internship__company=company).count()
    
    context = {
        'company': company,
        'internships': internships,
        'total_applications': total_applications,
    }
    return render(request, 'myapp/employer_dashboard.html', context)

@login_required
def create_company(request):
    user_profile, created = UserProfile.objects.get_or_create(user=request.user)
    if user_profile.user_type != 'employer':
        messages.error(request, 'Access denied. Employer account required.')
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = CompanyForm(request.POST, request.FILES)
        if form.is_valid():
            company = form.save()
            messages.success(request, 'Company profile created successfully!')
            return redirect('employer_dashboard')
    else:
        form = CompanyForm()
    
    return render(request, 'myapp/create_company.html', {'form': form})

@login_required
def create_internship(request):
    user_profile, created = UserProfile.objects.get_or_create(user=request.user)
    if user_profile.user_type != 'employer':
        messages.error(request, 'Access denied. Employer account required.')
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = InternshipForm(request.POST)
        if form.is_valid():
            internship = form.save(commit=False)
            # In real app, you'd get the company from user's profile
            company = Company.objects.filter(name__icontains=request.user.first_name).first()
            if company:
                internship.company = company
                internship.save()
                messages.success(request, 'Internship created successfully!')
                return redirect('employer_dashboard')
            else:
                messages.error(request, 'Please create a company profile first.')
    else:
        form = InternshipForm()
    
    return render(request, 'myapp/create_internship.html', {'form': form})

# Keep existing views for backward compatibility
@login_required(login_url='login')
def user_form(request):
    if request.method == 'POST':
        first_name = request.POST['first_name']
        last_name = request.POST['last_name']
        
        try:
            database_design_and_implementation = int(request.POST['database_design_and_implementation'])
            cloud_computing_and_techniques = int(request.POST['cloud_computing_and_techniques'])
            python_programming_methodology_2 = int(request.POST['python_programming_methodology_2'])
            internet_of_things_design_principles = int(request.POST['internet_of_things_design_principles'])
            big_data_architecture_and_programming = int(request.POST['big_data_architecture_and_programming'])
            network_technologies_and_design = int(request.POST['network_technologies_and_design'])

            if any(mark < 0 or mark > 100 for mark in [
                database_design_and_implementation, 
                cloud_computing_and_techniques, 
                python_programming_methodology_2, 
                internet_of_things_design_principles, 
                big_data_architecture_and_programming, 
                network_technologies_and_design]):
                return HttpResponse("Marks must be between 0 and 100.")

        except ValueError:
            return HttpResponse("Invalid input for marks. Please enter numeric values.")

        StudentRecord.objects.create(
            first_name=first_name,
            last_name=last_name,
            database_design_and_implementation=database_design_and_implementation,
            cloud_computing_and_techniques=cloud_computing_and_techniques,
            python_programming_methodology_2=python_programming_methodology_2,
            internet_of_things_design_principles=internet_of_things_design_principles,
            big_data_architecture_and_programming=big_data_architecture_and_programming,
            network_technologies_and_design=network_technologies_and_design
        )
        
        return redirect('view_records')  
    
    return render(request, 'myapp/form.html')

@login_required(login_url='login')
def view_records(request):
    search_query = request.GET.get('search', '')
    
    if search_query:
        search_terms = search_query.split()
        
        if len(search_terms) > 1:
            records = StudentRecord.objects.filter(
                Q(first_name__icontains=search_terms[0]) &
                Q(last_name__icontains=search_terms[-1])
            )
        else:
            records = StudentRecord.objects.filter(
                Q(first_name__icontains=search_query) |
                Q(last_name__icontains=search_query)
            )
    else:
        records = StudentRecord.objects.all()
    
    return render(request, 'myapp/view_records.html', {
        'records': records,
    })

@login_required(login_url='login')
def home(request):
    return render(request, 'myapp/home.html')

@login_required
def portal(request):
    primary_cv = (
        CV.objects.filter(user=request.user, is_primary=True).first()
        or CV.objects.filter(user=request.user).order_by('-uploaded_at').first()
    )
    return render(request, 'myapp/portal.html', {'primary_cv': primary_cv})

def download_jobs_excel(request):
    file_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'interntrack_jobs.xlsx')
    if os.path.exists(file_path):
        return FileResponse(open(file_path, 'rb'), as_attachment=True, filename='interntrack_jobs.xlsx')
    else:
        raise Http404("Excel file not found.")

@login_required
def download_internships_excel(request):
    if not request.user.userprofile.is_premium:
        messages.warning(request, 'Upgrade to premium to download analytics!')
        return redirect('dashboard')
    from .models import Internship
    qs = Internship.objects.all().select_related('company')
    data = []
    for i in qs:
        def tz_naive(dt):
            return dt.replace(tzinfo=None) if hasattr(dt, 'tzinfo') and dt.tzinfo else dt
        data.append({
            'Title': i.title,
            'Company': i.company.name if i.company else '',
            'Location': i.location,
            'Description': i.description,
            'Duration': i.duration,
            'Stipend': i.stipend,
            'Status': i.status,
            'Start Date': tz_naive(i.start_date) if i.start_date else '',
            'End Date': tz_naive(i.end_date) if i.end_date else '',
            'Application Deadline': tz_naive(i.application_deadline) if i.application_deadline else '',
            'Created': tz_naive(i.created_at) if hasattr(i, 'created_at') and i.created_at else '',
        })
    import pandas as pd
    df = pd.DataFrame(data)
    excel_path = 'all_internships.xlsx'
    df.to_excel(excel_path, index=False)
    return FileResponse(open(excel_path, 'rb'), as_attachment=True, filename='all_internships.xlsx')

@login_required
def confirm_premium_success(request):
    return payment_success(request)
          


@login_required
def application_studio(request, pk):
    job = get_object_or_404(Internship.objects.select_related('company'), pk=pk)
    profile, _ = UserProfile.objects.get_or_create(user=request.user)
    cvs = CV.objects.filter(user=request.user).order_by('-is_primary', '-uploaded_at')
    selected_cv = None
    cv_id = request.POST.get('cv_id') or request.GET.get('cv')
    if cv_id:
        selected_cv = get_object_or_404(cvs, pk=cv_id)
    else:
        selected_cv = cvs.first()
    if not selected_cv:
        messages.info(request, 'Upload a CV before opening the Application Studio.')
        return redirect('cv_list')
    pack = build_application_pack(user=request.user, cv=selected_cv, job=job)
    saved_documents = CareerDocument.objects.filter(user=request.user, internship=job, cv=selected_cv)
    if request.method == 'POST':
        action = request.POST.get('action')
        premium_actions = {'tailored_resume', 'cover_letter', 'application_email'}
        if action in premium_actions and not profile.is_premium:
            messages.warning(request, 'Document generation is a Premium feature. Your match analysis remains available below.')
            return redirect('application_studio', pk=job.pk)
        field_map = {
            'tailored_resume': ('Tailored resume', pack['tailored_resume']),
            'cover_letter': ('Cover letter', pack['cover_letter']),
            'application_email': ('Application email', pack['application_email']),
        }
        if action in field_map:
            label, default_content = field_map[action]
            content = request.POST.get('content', '').strip() or default_content
            doc = CareerDocument.objects.create(
                user=request.user, internship=job, cv=selected_cv,
                document_type=action, title=f'{label} — {job.title}', content=content,
            )
            messages.success(request, f'{label} saved. Review every statement before using it.')
            return redirect('career_document_edit', pk=doc.pk)
    return render(request, 'myapp/application_studio.html', {
        'internship': job, 'cvs': cvs, 'selected_cv': selected_cv, 'pack': pack,
        'saved_documents': saved_documents, 'is_premium': profile.is_premium,
    })


@login_required
def career_document_edit(request, pk):
    document = get_object_or_404(CareerDocument.objects.select_related('internship','cv'), pk=pk, user=request.user)
    if request.method == 'POST':
        document.title = request.POST.get('title', document.title).strip() or document.title
        document.content = request.POST.get('content', document.content).strip()
        document.save(update_fields=['title','content','updated_at'])
        messages.success(request, 'Document saved.')
        return redirect('career_document_edit', pk=document.pk)
    return render(request, 'myapp/career_document_edit.html', {'document': document})


@login_required
def career_document_download(request, pk):
    document = get_object_or_404(CareerDocument, pk=pk, user=request.user)
    from io import BytesIO
    from docx import Document
    safe_name = ''.join(c if c.isalnum() or c in '-_' else '_' for c in document.title)[:80]
    docx = Document()
    docx.add_heading(document.title, 0)
    for block in document.content.split('\n\n'):
        docx.add_paragraph(block)
    output = BytesIO(); docx.save(output); output.seek(0)
    response = HttpResponse(output.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}.docx"'
    return response
